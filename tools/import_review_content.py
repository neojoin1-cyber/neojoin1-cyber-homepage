#!/usr/bin/env python
"""Import a locked review manuscript into the linked Supabase project.

The source file never enters the public web tree. The importer extracts review
blocks locally, records the original SHA-256, applies one transaction through
the linked Supabase CLI, and deletes its temporary SQL file afterwards.
"""

from __future__ import annotations

import argparse
import base64
import hashlib
import json
import os
import re
import subprocess
import sys
import tempfile
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path
from typing import Any


MAX_SOURCE_BYTES = 100 * 1024 * 1024
MAX_TOTAL_TEXT = 25 * 1024 * 1024
MAX_BLOCKS = 5000
MAX_BLOCK_TEXT = 200_000
ALLOWED_PROGRAMS = {"civil", "elementary", "secondary"}
ALLOWED_STATUSES = {"draft", "review_ready"}


@dataclass
class Block:
    heading: str
    body: str


def clean_text(value: Any) -> str:
    text = str(value or "").replace("\x00", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def paragraph_blocks(text: str, prefix: str = "문단") -> list[Block]:
    parts = [clean_text(item) for item in re.split(r"\n\s*\n", clean_text(text))]
    return [Block(f"{prefix} {index}", item) for index, item in enumerate(parts, 1) if item]


def markdown_blocks(text: str) -> list[Block]:
    blocks: list[Block] = []
    heading = "원문"
    body: list[str] = []
    for line in clean_text(text).splitlines():
        match = re.match(r"^#{1,6}\s+(.+?)\s*$", line)
        if match:
            content = clean_text("\n".join(body))
            if content:
                blocks.append(Block(heading, content))
            heading = clean_text(match.group(1))
            body = []
        else:
            body.append(line)
    content = clean_text("\n".join(body))
    if content:
        blocks.append(Block(heading, content))
    return blocks or paragraph_blocks(text)


class ReviewHtmlParser(HTMLParser):
    SKIP = {"script", "style", "noscript", "template", "svg"}
    HEADINGS = {"h1", "h2", "h3", "h4", "h5", "h6"}
    CONTENT = {"p", "li", "blockquote", "pre", "td", "th"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.skip_depth = 0
        self.capture: str | None = None
        self.buffer: list[str] = []
        self.heading = "원문"
        self.blocks: list[Block] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag in self.SKIP:
            self.skip_depth += 1
        if not self.skip_depth and tag in self.HEADINGS | self.CONTENT:
            self.capture = tag
            self.buffer = []
        elif not self.skip_depth and tag == "br" and self.capture:
            self.buffer.append("\n")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in self.SKIP and self.skip_depth:
            self.skip_depth -= 1
            return
        if self.skip_depth or tag != self.capture:
            return
        text = clean_text("".join(self.buffer))
        if text and tag in self.HEADINGS:
            self.heading = text
        elif text:
            self.blocks.append(Block(self.heading, text))
        self.capture = None
        self.buffer = []

    def handle_data(self, data: str) -> None:
        if not self.skip_depth and self.capture:
            self.buffer.append(data)


def html_blocks(text: str) -> list[Block]:
    parser = ReviewHtmlParser()
    parser.feed(text)
    return parser.blocks or paragraph_blocks(re.sub(r"<[^>]+>", " ", text))


def json_blocks(value: Any) -> list[Block]:
    if isinstance(value, dict) and isinstance(value.get("blocks"), list):
        result = []
        for index, item in enumerate(value["blocks"], 1):
            if isinstance(item, dict):
                heading = clean_text(item.get("heading") or item.get("title") or f"문단 {index}")
                body = clean_text(item.get("body") or item.get("text") or item.get("content"))
                if body:
                    result.append(Block(heading, body))
        if result:
            return result
    if isinstance(value, list):
        result = []
        for index, item in enumerate(value, 1):
            if isinstance(item, dict):
                heading = clean_text(item.get("heading") or item.get("title") or f"항목 {index}")
                body = clean_text(item.get("body") or item.get("text") or item.get("content") or json.dumps(item, ensure_ascii=False, indent=2))
            else:
                heading, body = f"항목 {index}", clean_text(item)
            if body:
                result.append(Block(heading, body))
        if result:
            return result
    return [Block("JSON 원문", json.dumps(value, ensure_ascii=False, indent=2))]


def docx_blocks(path: Path) -> list[Block]:
    from docx import Document

    document = Document(path)
    blocks: list[Block] = []
    heading = "원문"
    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if not text:
            continue
        if paragraph.style and paragraph.style.name.lower().startswith("heading"):
            heading = text
        else:
            blocks.append(Block(heading, text))
    for table_index, table in enumerate(document.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            text = clean_text(" | ".join(cell.text for cell in row.cells))
            if text:
                blocks.append(Block(f"표 {table_index} · 행 {row_index}", text))
    return blocks


def pdf_blocks(path: Path) -> list[Block]:
    from pypdf import PdfReader

    reader = PdfReader(path)
    blocks = [Block(f"페이지 {index}", clean_text(page.extract_text() or "")) for index, page in enumerate(reader.pages, 1)]
    blocks = [block for block in blocks if block.body]
    if not blocks:
        raise ValueError("PDF에서 텍스트를 추출하지 못했습니다. OCR 적용본 또는 DOCX/HTML 원문이 필요합니다.")
    return blocks


def extract_blocks(path: Path) -> list[Block]:
    suffix = path.suffix.lower()
    if suffix in {".md", ".markdown"}:
        return markdown_blocks(path.read_text(encoding="utf-8-sig"))
    if suffix in {".txt", ".text"}:
        return paragraph_blocks(path.read_text(encoding="utf-8-sig"))
    if suffix in {".html", ".htm"}:
        return html_blocks(path.read_text(encoding="utf-8-sig"))
    if suffix == ".json":
        return json_blocks(json.loads(path.read_text(encoding="utf-8-sig")))
    if suffix == ".docx":
        return docx_blocks(path)
    if suffix == ".pdf":
        return pdf_blocks(path)
    raise ValueError(f"지원하지 않는 원문 형식입니다: {suffix or '(확장자 없음)'}")


def b64(value: str) -> str:
    return base64.b64encode(value.encode("utf-8")).decode("ascii")


def sql_text(value: str) -> str:
    return f"convert_from(decode('{b64(value)}', 'base64'), 'UTF8')"


def build_sql(manifest: dict[str, Any], source: Path, digest: str, blocks: list[Block]) -> str:
    program = manifest["program"]
    subject_code = manifest["subjectCode"]
    title = manifest["title"]
    version = manifest["version"]
    kind = manifest["kind"]
    stage = manifest.get("reviewStage", "1차 검수")
    status = manifest.get("status", "review_ready")
    rows = []
    for index, block in enumerate(blocks, 1):
        fingerprint = hashlib.sha256(block.body.encode("utf-8")).hexdigest()
        rows.append(
            f"(v_document, 'b{index:05d}', {sql_text(block.heading)}, {sql_text(block.body)}, {index}, '{fingerprint}')"
        )
    return f"""begin;
do $review_import$
declare
  v_subject uuid;
  v_document uuid;
begin
  select id into v_subject
    from public.review_subjects
   where program_id = {sql_text(program)} and code = {sql_text(subject_code)};
  if v_subject is null then
    raise exception '검수 과목을 찾지 못했습니다: %/%', {sql_text(program)}, {sql_text(subject_code)};
  end if;
  if exists (
    select 1 from public.review_documents
     where subject_id = v_subject and title = {sql_text(title)} and version = {sql_text(version)}
  ) then
    raise exception '동일 과목·제목·버전의 검수 자료가 이미 있습니다.';
  end if;
  insert into public.review_documents
    (subject_id, kind, title, version, review_stage, source_object_path, source_sha256, status)
  values
    (v_subject, {sql_text(kind)}, {sql_text(title)}, {sql_text(version)}, {sql_text(stage)}, null, '{digest}', {sql_text(status)})
  returning id into v_document;
  insert into public.review_blocks
    (document_id, block_key, heading, body, sort_order, source_fingerprint)
  values
    {',\n    '.join(rows)};
  raise notice 'review_document_id=% blocks={len(blocks)} sha256={digest}', v_document;
end
$review_import$;
commit;
"""


def validate_manifest(manifest: dict[str, Any]) -> None:
    required = ["program", "subjectCode", "kind", "title", "version", "source"]
    missing = [key for key in required if not clean_text(manifest.get(key))]
    if missing:
        raise ValueError(f"manifest 필수 항목 누락: {', '.join(missing)}")
    if manifest["program"] not in ALLOWED_PROGRAMS:
        raise ValueError("program은 civil, elementary, secondary 중 하나여야 합니다.")
    if manifest.get("status", "review_ready") not in ALLOWED_STATUSES:
        raise ValueError("status는 draft 또는 review_ready만 허용됩니다.")


def main() -> int:
    parser = argparse.ArgumentParser(description="설탕과소금 전문위원 검수 원문 적재")
    parser.add_argument("manifest", type=Path, help="원문 메타데이터 JSON")
    parser.add_argument("--apply", action="store_true", help="연결된 Supabase 운영 DB에 실제 반영")
    args = parser.parse_args()

    manifest_path = args.manifest.resolve()
    manifest = json.loads(manifest_path.read_text(encoding="utf-8-sig"))
    validate_manifest(manifest)
    source = Path(manifest["source"])
    if not source.is_absolute():
        source = (manifest_path.parent / source).resolve()
    if not source.is_file():
        raise FileNotFoundError(f"원문 파일을 찾지 못했습니다: {source}")
    if source.stat().st_size > MAX_SOURCE_BYTES:
        raise ValueError("원문 파일이 100MB를 초과합니다.")

    source_bytes = source.read_bytes()
    digest = hashlib.sha256(source_bytes).hexdigest()
    blocks = [Block(clean_text(block.heading), clean_text(block.body)) for block in extract_blocks(source)]
    blocks = [block for block in blocks if block.heading and block.body]
    if not blocks:
        raise ValueError("검수 문단을 한 건도 추출하지 못했습니다.")
    if len(blocks) > MAX_BLOCKS:
        raise ValueError(f"검수 문단이 {MAX_BLOCKS}건을 초과합니다.")
    if any(len(block.body) > MAX_BLOCK_TEXT for block in blocks):
        raise ValueError("단일 검수 문단이 200,000자를 초과합니다. 원문을 더 작은 장·절로 나눠 주세요.")
    total_text = sum(len(block.heading) + len(block.body) for block in blocks)
    if total_text > MAX_TOTAL_TEXT:
        raise ValueError("추출된 원문 텍스트가 25MB를 초과합니다.")

    print(json.dumps({
        "status": "ready_to_apply" if args.apply else "validated",
        "source": str(source),
        "sha256": digest,
        "blocks": len(blocks),
        "characters": total_text,
        "program": manifest["program"],
        "subjectCode": manifest["subjectCode"],
        "title": manifest["title"],
        "version": manifest["version"],
    }, ensure_ascii=False, indent=2))
    if not args.apply:
        return 0

    sql = build_sql(manifest, source, digest, blocks)
    temp_path: str | None = None
    try:
        with tempfile.NamedTemporaryFile("w", encoding="utf-8", suffix=".sql", delete=False) as temp:
            temp.write(sql)
            temp_path = temp.name
        completed = subprocess.run(
            ["supabase", "db", "query", "--linked", "--file", temp_path],
            check=False,
            text=True,
            capture_output=True,
        )
        if completed.returncode != 0:
            print(completed.stderr or completed.stdout, file=sys.stderr)
            return completed.returncode
        print("IMPORT_APPLIED")
        return 0
    finally:
        if temp_path:
            try:
                os.remove(temp_path)
            except FileNotFoundError:
                pass


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as error:
        print(f"IMPORT_FAILED: {error}", file=sys.stderr)
        raise SystemExit(1)
